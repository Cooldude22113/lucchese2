"""
documents/generation/docx_generator.py
Convert markdown-ish content to a styled .docx file (Rule 12).

Handles # ## ### headings, plain "Section:" headers, - / * bullets, 1. numbered
lists, **bold**, *italic*, and --- rules. Saves under data/generated_docs/ and
returns the filepath.
"""

from __future__ import annotations

import re
import uuid
from datetime import datetime, timezone

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from config.paths import GENERATED_DOCS_DIR
from documents.generation import document_styles as S
from documents.generation.inline_formatting import add_inline_formatting
from documents.generation.markdown_parser import looks_like_section_header


def _rgb(color: tuple[int, int, int]) -> RGBColor:
    return RGBColor(*color)


def markdown_to_docx(content: str, title: str) -> str:
    """Render `content` to a .docx titled `title`; returns the saved filepath."""
    doc = DocxDocument()

    # Page margins.
    for section in doc.sections:
        section.top_margin = Inches(S.MARGIN_TOP)
        section.bottom_margin = Inches(S.MARGIN_BOTTOM)
        section.left_margin = Inches(S.MARGIN_LEFT)
        section.right_margin = Inches(S.MARGIN_RIGHT)

    # Title.
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run(title.replace("_", " ").title())
    run.bold = True
    run.font.size = Pt(S.SIZE_TITLE)
    run.font.color.rgb = _rgb(S.INK)

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(datetime.now(timezone.utc).strftime("Generated %d %B %Y"))
    date_run.font.size = Pt(S.SIZE_DATE)
    date_run.font.color.rgb = _rgb(S.MUTED)
    doc.add_paragraph()  # spacer

    # Body.
    for line in content.split("\n"):
        s = line.strip()

        if not s:
            doc.add_paragraph()
            continue

        if s.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(s[2:].strip())
            run.bold = True
            run.font.size = Pt(S.SIZE_H1)
            run.font.color.rgb = _rgb(S.INK)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)

        elif s.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(s[3:].strip())
            run.bold = True
            run.font.size = Pt(S.SIZE_H2)
            run.font.color.rgb = _rgb(S.INK)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(3)

        elif s.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(s[4:].strip())
            run.bold = True
            run.font.size = Pt(S.SIZE_H3)
            run.font.color.rgb = _rgb(S.SUBHEAD)
            p.paragraph_format.space_before = Pt(8)

        elif looks_like_section_header(s):
            p = doc.add_paragraph()
            run = p.add_run(s.rstrip(":"))
            run.bold = True
            run.font.size = Pt(S.SIZE_SECTION)
            run.font.color.rgb = _rgb(S.INK)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(3)

        elif s.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_formatting(p, s[2:].strip())

        elif re.match(r"^\d+\.\s", s):
            p = doc.add_paragraph(style="List Number")
            add_inline_formatting(p, re.sub(r"^\d+\.\s*", "", s))

        elif s in ("---", "***", "___"):
            p = doc.add_paragraph()
            run = p.add_run("─" * 60)
            run.font.color.rgb = _rgb(S.RULE)
            run.font.size = Pt(S.SIZE_RULE)

        else:
            p = doc.add_paragraph()
            add_inline_formatting(p, s)
            p.paragraph_format.space_after = Pt(4)

    # Save.
    GENERATED_DOCS_DIR.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r"[^\w\-]", "_", title)[:60]
    filename = f"{safe_name}_{uuid.uuid4().hex[:6]}.docx"
    filepath = GENERATED_DOCS_DIR / filename
    doc.save(str(filepath))
    return str(filepath)
